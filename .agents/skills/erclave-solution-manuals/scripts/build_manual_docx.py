from argparse import ArgumentParser
from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK


def build(source: Path, target: Path) -> None:
    document = Document()
    for raw in source.read_text(encoding="utf-8").splitlines():
        line = raw.strip()
        if not line:
            continue
        if line == "---page---":
            document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
        elif line.startswith("### "):
            document.add_heading(line[4:], level=3)
        elif line.startswith("## "):
            document.add_heading(line[3:], level=2)
        elif line.startswith("# "):
            document.add_heading(line[2:], level=1)
        elif line.startswith("- "):
            document.add_paragraph(line[2:], style="List Bullet")
        elif len(line) > 3 and line[0].isdigit() and line[1:3] == ". ":
            document.add_paragraph(line[3:], style="List Number")
        else:
            document.add_paragraph(line)
    target.parent.mkdir(parents=True, exist_ok=True)
    document.save(target)


def main() -> None:
    parser = ArgumentParser(description="Genera un manual DOCX desde su fuente Markdown.")
    parser.add_argument("source", type=Path)
    parser.add_argument("target", type=Path)
    args = parser.parse_args()
    build(args.source, args.target)


if __name__ == "__main__":
    main()
