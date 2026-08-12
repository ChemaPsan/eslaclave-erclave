"""Generate the ERClave QA Word guide from its Markdown source."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "qa" / "guia_pruebas_qa_mvp.md"
OUTPUT = ROOT / "docs" / "qa" / "guia_pruebas_qa_mvp.docx"

PURPLE = "5B2A86"
PURPLE_DARK = "351455"
LAVENDER = "EDE2F5"
WHITE = "FFFFFF"
GRAY = "F3F3F3"


def clean(value: str) -> str:
    value = re.sub(r"`([^`]+)`", r"\1", value)
    value = value.replace("**", "")
    return value.strip()


def shade(cell, color: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    fill = properties.find(qn("w:shd"))
    if fill is None:
        fill = OxmlElement("w:shd")
        properties.append(fill)
    fill.set(qn("w:fill"), color)


def set_cell_margins(cell, top=80, start=100, bottom=80, end=100) -> None:
    properties = cell._tc.get_or_add_tcPr()
    margins = properties.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        properties.append(margins)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    properties = row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    properties.append(repeat)


def add_toc(document: Document) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = 'TOC \\o "1-3" \\h \\z \\u'
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "Actualiza este campo en Word para mostrar el índice."
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for element in (begin, instruction, separate, placeholder, end):
        run._r.append(element)


def add_inline_markdown(paragraph, text: str) -> None:
    parts = re.split(r"(\*\*.*?\*\*|`.*?`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            paragraph.add_run(part[2:-2]).bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor.from_string(PURPLE_DARK)
        else:
            paragraph.add_run(part)


def configure_styles(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.08

    for name, size, color in [
        ("Title", 26, PURPLE_DARK),
        ("Heading 1", 18, PURPLE_DARK),
        ("Heading 2", 14, PURPLE),
        ("Heading 3", 11.5, PURPLE),
    ]:
        style = document.styles[name]
        style.font.name = "Aptos Display"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)

    if "QA Note" not in document.styles:
        note = document.styles.add_style("QA Note", WD_STYLE_TYPE.PARAGRAPH)
        note.base_style = document.styles["Normal"]
        note.font.color.rgb = RGBColor.from_string(PURPLE_DARK)
        note.paragraph_format.left_indent = Cm(0.6)
        note.paragraph_format.right_indent = Cm(0.6)
        note.paragraph_format.space_before = Pt(5)
        note.paragraph_format.space_after = Pt(8)


def configure_page(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

    header = section.header.paragraphs[0]
    header.text = "ERClave · Guía QA del MVP"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.runs[0].font.size = Pt(8)
    header.runs[0].font.color.rgb = RGBColor.from_string(PURPLE)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("Documento de trabajo QA · No incluir credenciales ni secretos · Página ")
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    footer._p.append(field)
    for run in footer.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(100, 100, 100)


def add_cover(document: Document) -> None:
    document.add_paragraph()
    document.add_paragraph()
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("ERClave")
    title_run.bold = True
    title_run.font.size = Pt(34)
    title_run.font.color.rgb = RGBColor.from_string(PURPLE_DARK)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Guía de pruebas QA del MVP")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor.from_string(PURPLE)

    description = document.add_paragraph()
    description.alignment = WD_ALIGN_PARAGRAPH.CENTER
    description.add_run(
        "Qué probar · Cómo hacerlo · Resultado esperado · Valor para el negocio"
    ).italic = True

    document.add_paragraph()
    metadata = document.add_table(rows=4, cols=2)
    metadata.alignment = WD_TABLE_ALIGNMENT.CENTER
    values = [
        ("Ambiente", "QA"),
        ("Aplicación", "https://erclave.web.app"),
        ("Backoffice", "https://erclave.web.app/backoffice/"),
        ("Actualización", "24 de julio de 2026"),
    ]
    for row, (label, value) in zip(metadata.rows, values):
        row.cells[0].text = label
        row.cells[1].text = value
        row.cells[0].paragraphs[0].runs[0].bold = True
        shade(row.cells[0], LAVENDER)
        for cell in row.cells:
            set_cell_margins(cell)

    document.add_paragraph()
    warning = document.add_paragraph(style="QA Note")
    warning.alignment = WD_ALIGN_PARAGRAPH.CENTER
    warning.add_run(
        "Uso interno de QA. Nunca registrar contraseñas, tokens o ligas vigentes de invitación."
    ).bold = True
    document.add_page_break()


def add_table(document: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    width = max(len(row) for row in rows)
    table = document.add_table(rows=0, cols=width)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for row_index, values in enumerate(rows):
        row = table.add_row()
        for column in range(width):
            value = values[column] if column < len(values) else ""
            cell = row.cells[column]
            cell.text = clean(value)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(2)
                for run in paragraph.runs:
                    run.font.size = Pt(8.5)
            if row_index == 0:
                shade(cell, PURPLE_DARK)
                for run in cell.paragraphs[0].runs:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor.from_string(WHITE)
            elif row_index % 2 == 0:
                shade(cell, GRAY)
        if row_index == 0:
            set_repeat_table_header(row)
    document.add_paragraph()


def render_markdown(document: Document, text: str) -> None:
    lines = text.splitlines()
    index = 0
    in_code = False
    code_lines: list[str] = []
    while index < len(lines):
        line = lines[index]

        if line.startswith("```"):
            if in_code:
                paragraph = document.add_paragraph()
                paragraph.style = "QA Note"
                run = paragraph.add_run("\n".join(code_lines))
                run.font.name = "Consolas"
                run.font.size = Pt(9)
                code_lines = []
                in_code = False
            else:
                in_code = True
            index += 1
            continue
        if in_code:
            code_lines.append(line)
            index += 1
            continue

        if line.startswith("|"):
            table_rows = []
            while index < len(lines) and lines[index].startswith("|"):
                if "---" not in lines[index]:
                    table_rows.append([cell.strip() for cell in lines[index].strip("|").split("|")])
                index += 1
            add_table(document, table_rows)
            continue

        heading = re.match(r"^(#{1,3})\s+(.+)$", line)
        if heading:
            level = len(heading.group(1))
            title = clean(heading.group(2))
            if level == 1:
                index += 1
                continue
            document.add_heading(title, level=level - 1)
            index += 1
            continue

        if re.match(r"^\d+\.\s+", line):
            paragraph = document.add_paragraph(style="List Number")
            add_inline_markdown(paragraph, re.sub(r"^\d+\.\s+", "", line))
        elif line.startswith("- "):
            paragraph = document.add_paragraph(style="List Bullet")
            add_inline_markdown(paragraph, line[2:])
        elif line.startswith("> "):
            paragraph = document.add_paragraph(style="QA Note")
            add_inline_markdown(paragraph, line[2:])
        elif line.strip():
            paragraph = document.add_paragraph()
            add_inline_markdown(paragraph, line.strip())
        index += 1


def build() -> None:
    source = SOURCE.read_text(encoding="utf-8")
    document = Document()
    configure_styles(document)
    configure_page(document)
    add_cover(document)
    document.add_heading("Índice", level=1)
    add_toc(document)
    document.add_page_break()
    render_markdown(document, source)

    properties = document.core_properties
    properties.title = "ERClave - Guía de pruebas QA del MVP"
    properties.subject = "Casos manuales de prueba, contexto y evidencia"
    properties.author = "ERClave / Codex"
    properties.keywords = "ERClave, QA, MVP, pruebas, multitenant"
    document.save(OUTPUT)


def validate() -> None:
    document = Document(OUTPUT)
    full_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    assert "Guía de pruebas QA del MVP" in full_text
    assert "QA-BO-01" in full_text
    assert "QA-PROD-04" in full_text
    assert "Criterio de salida" in full_text
    assert len(document.tables) >= 7
    assert OUTPUT.stat().st_size > 20_000


if __name__ == "__main__":
    build()
    validate()
    print(
        f"Created {OUTPUT.relative_to(ROOT)}: "
        f"{OUTPUT.stat().st_size} bytes, validated with python-docx."
    )
